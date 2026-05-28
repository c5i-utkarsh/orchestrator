"""File discovery and loading for txt, md, json (and optionally pdf)."""

from __future__ import annotations

import hashlib
import json
import logging
import uuid
from pathlib import Path
from typing import Iterator

from .models import Document, RejectionReason

logger = logging.getLogger(__name__)

_PLAIN_TEXT_EXTENSIONS = {".txt", ".md"}


def load_corpus(
    input_dir: str,
    supported_extensions: list[str],
    enable_pdf: bool = False,
) -> list[Document]:
    """Walk *input_dir* and load every supported file into a Document list.

    Args:
        input_dir: Root directory to scan recursively.
        supported_extensions: List of extensions to include (e.g. [".txt", ".md"]).
        enable_pdf: When True, also attempt to load .pdf files via PyMuPDF.

    Returns:
        List of Document objects; failed loads are included with accepted=False.
    """
    root = Path(input_dir)
    if not root.exists():
        raise FileNotFoundError(f"Input directory not found: {root.resolve()}")

    extensions = set(supported_extensions)
    if enable_pdf:
        extensions.add(".pdf")

    docs: list[Document] = []
    for path in _walk(root, extensions):
        doc = _load_single(path)
        docs.append(doc)

    logger.info("Loaded %d documents from '%s'", len(docs), root)
    return docs


def _walk(root: Path, extensions: set[str]) -> Iterator[Path]:
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in extensions:
            yield path


def _load_single(path: Path) -> Document:
    doc_id = str(uuid.uuid4())
    ext = path.suffix.lower()
    try:
        raw_text = _read(path, ext)
        content_hash = _sha256(raw_text)
        return Document(
            doc_id=doc_id,
            source_path=str(path),
            extension=ext,
            raw_text=raw_text,
            content_hash=content_hash,
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to load '%s': %s", path, exc)
        return Document(
            doc_id=doc_id,
            source_path=str(path),
            extension=ext,
            accepted=False,
            rejection_reason=RejectionReason.LOAD_ERROR,
            rejection_detail=str(exc),
        )


def _read(path: Path, ext: str) -> str:
    if ext in _PLAIN_TEXT_EXTENSIONS:
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".json":
        return _read_json(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".xlsx":
        return _read_xlsx(path)
    if ext == ".csv":
        return _read_csv(path)
    raise ValueError(f"Unsupported extension: {ext}")


def _read_json(path: Path) -> str:
    """Extract text values from a JSON file (handles dict or list at root)."""
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    if isinstance(data, dict):
        return _flatten_dict(data)
    if isinstance(data, list):
        return "\n".join(_flatten_dict(item) if isinstance(item, dict) else str(item) for item in data)
    return str(data)


def _flatten_dict(d: dict) -> str:
    parts: list[str] = []
    for v in d.values():
        if isinstance(v, str):
            parts.append(v)
        elif isinstance(v, (dict, list)):
            parts.append(_flatten_dict(v) if isinstance(v, dict) else " ".join(str(x) for x in v))
    return " ".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError("PyMuPDF is required for PDF support: pip install PyMuPDF") from exc

    pages: list[str] = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            pages.append(page.get_text())
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document as _DocxDoc
    except ImportError as exc:
        raise ImportError("python-docx required: pip install python-docx") from exc
    doc = _DocxDoc(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _read_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError as exc:
        raise ImportError("openpyxl required: pip install openpyxl") from exc
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            vals = [str(c) for c in row if c is not None and str(c).strip()]
            if vals:
                lines.append("\t".join(vals))
    wb.close()
    return "\n".join(lines)


def _read_csv(path: Path) -> str:
    import csv
    rows: list[str] = []
    with path.open(encoding="utf-8", errors="replace", newline="") as f:
        for row in csv.reader(f):
            if any(c.strip() for c in row):
                rows.append("\t".join(row))
    return "\n".join(rows)


def _sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()
