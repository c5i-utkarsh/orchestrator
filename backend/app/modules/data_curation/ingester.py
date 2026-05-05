import uuid
import json
from pathlib import Path
from dataclasses import dataclass, field

from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings

settings = get_settings()


@dataclass
class CanonicalDocument:
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    text: str = ""
    title: str = ""
    source: str = ""
    metadata: dict = field(default_factory=dict)


class Ingester:
    """
    Normalizes all input types (files, URLs, DB exports) into
    a list of CanonicalDocument objects.
    """

    SUPPORTED_EXTENSIONS = {".txt", ".csv", ".json", ".jsonl", ".parquet",
                             ".pdf", ".docx", ".md", ".rst", ".html"}

    async def ingest_directory(self, corpus_dir: str) -> list[CanonicalDocument]:
        docs = []
        corpus_path = Path(corpus_dir)
        for file_path in corpus_path.rglob("*"):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                file_docs = await self._ingest_file(file_path)
                docs.extend(file_docs)
        return docs

    async def _ingest_file(self, path: Path) -> list[CanonicalDocument]:
        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                return self._read_pdf(path)
            elif suffix == ".docx":
                return self._read_docx(path)
            elif suffix in (".json", ".jsonl"):
                return self._read_json(path)
            elif suffix == ".parquet":
                return self._read_parquet(path)
            elif suffix == ".csv":
                return self._read_csv(path)
            else:
                return self._read_text(path)
        except Exception as e:
            return [CanonicalDocument(
                title=path.name,
                text=f"[ERROR reading {path.name}: {e}]",
                source=str(path),
            )]

    def _read_pdf(self, path: Path) -> list[CanonicalDocument]:
        import pdfplumber
        docs = []
        with pdfplumber.open(path) as pdf:
            full_text = "\n".join(
                page.extract_text() or "" for page in pdf.pages
            )
        docs.append(CanonicalDocument(
            title=path.stem,
            text=full_text.strip(),
            source=str(path),
            metadata={"type": "pdf", "pages": len(pdf.pages) if hasattr(pdf, "pages") else 0},
        ))
        return docs

    def _read_docx(self, path: Path) -> list[CanonicalDocument]:
        from docx import Document
        doc = Document(str(path))
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        return [CanonicalDocument(
            title=path.stem, text=text, source=str(path),
            metadata={"type": "docx"},
        )]

    def _read_json(self, path: Path) -> list[CanonicalDocument]:
        docs = []
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            if path.suffix == ".jsonl":
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        obj = json.loads(line)
                        text = obj.get("text") or obj.get("content") or str(obj)
                        docs.append(CanonicalDocument(
                            title=obj.get("title", path.stem),
                            text=text, source=str(path),
                            metadata={"type": "jsonl"},
                        ))
                    except Exception:
                        continue
            else:
                data = json.load(f)
                if isinstance(data, list):
                    for item in data:
                        text = item.get("text") or item.get("content") or str(item)
                        docs.append(CanonicalDocument(
                            title=item.get("title", path.stem),
                            text=text, source=str(path),
                            metadata={"type": "json"},
                        ))
                else:
                    docs.append(CanonicalDocument(
                        title=path.stem, text=str(data),
                        source=str(path), metadata={"type": "json"},
                    ))
        return docs

    def _read_parquet(self, path: Path) -> list[CanonicalDocument]:
        import pandas as pd
        df = pd.read_parquet(path)
        text_col = next(
            (c for c in ["text", "content", "body", "description"] if c in df.columns),
            df.columns[0]
        )
        docs = []
        for _, row in df.iterrows():
            docs.append(CanonicalDocument(
                title=str(row.get("title", path.stem)),
                text=str(row[text_col]),
                source=str(path),
                metadata={"type": "parquet"},
            ))
        return docs

    def _read_csv(self, path: Path) -> list[CanonicalDocument]:
        import pandas as pd
        df = pd.read_csv(path)
        text_col = next(
            (c for c in ["text", "content", "body", "description"] if c in df.columns),
            df.columns[0]
        )
        docs = []
        for _, row in df.iterrows():
            text = " | ".join(f"{col}: {row[col]}" for col in df.columns
                              if str(row[col]).strip())
            docs.append(CanonicalDocument(
                title=str(row.get("title", path.stem)),
                text=text, source=str(path),
                metadata={"type": "csv"},
            ))
        return docs

    def _read_text(self, path: Path) -> list[CanonicalDocument]:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [CanonicalDocument(
            title=path.stem, text=text.strip(),
            source=str(path), metadata={"type": path.suffix},
        )]

    def save_to_corpus_dir(
        self, docs: list[CanonicalDocument], corpus_dir: str
    ) -> str:
        """Save canonical docs as JSONL for graphifyy consumption."""
        out_path = Path(corpus_dir) / "canonical_corpus.jsonl"
        out_path.parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as f:
            for doc in docs:
                f.write(json.dumps({
                    "id": doc.id,
                    "title": doc.title,
                    "text": doc.text,
                    "source": doc.source,
                    "metadata": doc.metadata,
                }) + "\n")
        return str(out_path)
